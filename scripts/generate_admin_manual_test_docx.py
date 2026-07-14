"""Generate ScanLink Admin Manual Test Guide DOCX."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def set_run(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=11)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run(run, size=11)
    return p


def step(num, title, does, how, expect):
    h3(f"Step {num}: {title}")

    p = doc.add_paragraph()
    r = p.add_run("What it does: ")
    set_run(r, bold=True, size=11)
    r2 = p.add_run(does)
    set_run(r2, size=11)

    p2 = doc.add_paragraph()
    r = p2.add_run("How to test: ")
    set_run(r, bold=True, size=11)
    r2 = p2.add_run(how)
    set_run(r2, size=11)

    p3 = doc.add_paragraph()
    r = p3.add_run("Expected result: ")
    set_run(r, bold=True, size=11)
    r2 = p3.add_run(expect)
    set_run(r2, size=11)

    p4 = doc.add_paragraph()
    r = p4.add_run("Pass / Fail: ___________    Notes: _______________________________")
    set_run(r, size=10, color=(89, 89, 89))


# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ScanLink Admin Panel")
set_run(r, bold=True, size=22, color=(15, 76, 129))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Full Manual Test Guide")
set_run(r, bold=True, size=16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Laravel Filament Admin  ·  Replaces legacy siteadmin  ·  Dated 13 Jul 2026"
)
set_run(r, size=10, color=(102, 102, 102))

doc.add_paragraph()

h1("1. Purpose")
para(
    "This document is a complete manual test script for the ScanLink Admin Panel at "
    "http://localhost:8000/admin. Follow every step in order (or by section). "
    "Each step explains what the feature does, how to test it, and what you should see."
)
para(
    "Scope: Admin portal only (clients, users, profiles, orders, settings, CMS). "
    "Out of scope: client portal login, public scan pages, form-builder question editor, mobile apps."
)

h1("2. Before you start")
h2("2.1 Prerequisites")
bullet("Docker stack running for scanlink-laravel (app, mysql, redis).")
bullet("App reachable at http://localhost:8000")
bullet("Browser: Chrome or Edge recommended.")
bullet("Have a notepad ready to record Pass/Fail next to each step.")

h2("2.2 Start the app (if needed)")
bullet(r"Open a terminal in: D:\projects\product\scanlink-laravel")
bullet("Run: docker compose up -d")
bullet("Wait until healthy, then open: http://localhost:8000/admin/login")

h2("2.3 Test accounts")
bullet("Super Admin — Email: admin@scanlink.com  |  Password: Admin@12345")
bullet(
    "Support role (optional) — use a Support user to verify read-only restrictions."
)

h2("2.4 Legend")
bullet("What it does — business purpose of the screen/action.")
bullet("How to test — exact clicks and data to enter.")
bullet("Expected result — what success looks like.")
bullet("Pass / Fail — mark while testing.")

# AUTH
h1("3. Authentication")

step(
    1,
    "Open login page",
    "Loads the admin sign-in screen and blocks unauthenticated access to /admin.",
    "Open http://localhost:8000/admin (or /admin/login) while logged out.",
    "You see the Filament login form (email + password). Visiting /admin/clients redirects to login.",
)

step(
    2,
    "Invalid login",
    "Rejects bad credentials so unknown users cannot enter the admin panel.",
    "Enter a wrong email/password and click Sign in.",
    "You stay on the login page; an error or notification appears. You are not taken to Admin Home.",
)

step(
    3,
    "Valid Super Admin login",
    "Authenticates a platform admin and opens the control panel.",
    "Sign in with admin@scanlink.com / Admin@12345.",
    "You land on Admin Home (tile dashboard). Sidebar shows groups: Website, Order, Clients, Settings.",
)

step(
    4,
    "Forgot password page",
    "Lets an admin request a password-reset email link.",
    "Sign out (or use a private window). Open http://localhost:8000/admin/password-reset/request. "
    "Enter admin@scanlink.com and submit.",
    "Page accepts the request without a 500 error (email may go to Mailpit if configured).",
)

step(
    5,
    "Register page availability",
    "Confirms the Sign up page loads (registration may be used for new admin users depending on env).",
    "Open http://localhost:8000/admin/register while logged out.",
    "Sign up form loads successfully (no Whoops/500).",
)

# HOME
h1("4. Admin Home")

step(
    6,
    "Admin Home tiles",
    "Provides one-click shortcuts to every major admin area (legacy control-panel style).",
    "From Admin Home, click each of the 12 tiles in turn, then use Back or the sidebar to return.",
    "Each tile opens the correct page: Add Client, Manage Client, Sub Divide Client, "
    "Manage Profiles(Product), Manage Order, Manage Code Orders, Manage Form Builder Orders, "
    "Global Settings, Code Pricing, Manage Testimonial, Manage Gallery, Change Password.",
)

step(
    7,
    "Sidebar navigation",
    "Groups admin features the same way as legacy siteadmin menus.",
    "Expand Website, Order, Clients, Settings in the left sidebar and open one item from each group.",
    "Pages load without errors. Active nav item is highlighted.",
)

# CLIENTS
h1("5. Manage Client")

step(
    8,
    "Client list",
    "Shows all clients with search/filter and row actions (edit, block, users, etc.).",
    "Open Website → Manage Client (or tile Manage Client).",
    "Table lists clients. Filter/search by client name works. Reseller badge appears where applicable.",
)

step(
    9,
    "Create client — validation",
    "Prevents saving incomplete client records.",
    "Click Add Client (or tile). Leave required fields empty and click Create.",
    "Validation errors appear on required fields (client name, email, URL slug, etc.). No new client is created.",
)

step(
    10,
    "Create client — invalid URL slug",
    "Enforces URL character rules so public/portal paths stay clean.",
    'On Add Client, enter an invalid URL such as "bad url!!" in the URL field and try to save.',
    "URL field shows a validation error; save is blocked.",
)

step(
    11,
    "Create client — happy path (+ optional user)",
    "Creates a client company record and optionally the primary portal user on the same screen "
    "(legacy Add Client + Add User).",
    "Fill valid Client Name, Contact, Email, Telephone, Address, URL slug (letters/numbers/hyphen only). "
    "Optionally fill Add User email/password. Enable Checklist / CustomQR toggles if shown. Click Create.",
    "Success notification. Client appears in Manage Client list. If Add User was filled, that user "
    "appears under Manage User for the client.",
)

step(
    12,
    "Edit client",
    "Updates client details, reseller/free-code options, and feature toggles.",
    "Open an existing client → Edit. Change contact person or telephone. "
    "Toggle checklist_option / customqr_option. Save.",
    "Changes persist after refresh. Reseller name/code fields display correctly when reseller applies.",
)

step(
    13,
    "Block / unblock client",
    "Disables or re-enables a client account without deleting history.",
    "From Manage Client list, use Block on a test client, then Unblock.",
    "Status flips; blocked client is not treated as active. Unblock restores access flag.",
)

step(
    14,
    "Delete client (careful — use throwaway only)",
    "Removes a client from the manage list (use a test client created for this purpose only).",
    "On a throwaway test client, choose Delete and confirm.",
    "Client disappears from the active list (or is marked deleted per app rules).",
)

# USERS
h1("6. Manage User (per client)")

step(
    15,
    "Open users for a client",
    "Lists portal users belonging to one client (legacy user/index/{clientId}).",
    "From Manage Client, open Users for a client (or go to /admin/clients/{id}/users).",
    "Manage User page appears. Existing users are listed with email and expiry.",
)

step(
    16,
    "Add user",
    "Creates an additional portal user under the selected client.",
    "Click Add User. Enter email + password (and any other required fields). Save.",
    "New user row appears in the table for that client.",
)

step(
    17,
    "Inline expiry / renew account",
    "Extends or edits the user’s expiry date and can trigger renewal notification behaviour.",
    "Edit the Expires column (or renew action) on a user. Set a future date and save.",
    "Expiry updates in the list. No 500 error. (Email may be sent depending on mail config.)",
)

# SUBDIVIDE
h1("7. Sub Divide Client")

step(
    18,
    "Subdivide wizard",
    "Splits selected profiles from one client into a new or existing client (legacy multi-step subdivide).",
    "Open Clients → Sub Divide Client (or home tile). Complete wizard steps: choose source client → "
    "select profiles → choose destination / new client details → confirm.",
    "Wizard advances without errors. On finish, selected profiles belong to the destination client. "
    "Success notification appears.",
)

# PROFILES
h1("8. Manage Profiles (Product)")

step(
    19,
    "Profiles list",
    "Lists non-archived profiles with legacy columns (Profile No., Name, Note, Contact, Telephone).",
    "Open Website → Manage Profiles(Product).",
    "Table loads. Name filter and Type filter work. Archived/deleted profiles are hidden.",
)

step(
    20,
    "Add Profile — type picker",
    "Starts profile creation with client + equipment type (plant, location, asset, product, procedure, "
    "misc, customqr, code). People is excluded like legacy.",
    "Click Add Profile. Open Select Profile Type and Select Client.",
    "Types include CustomQR and Code; People is not listed. Both fields are required.",
)

step(
    21,
    "Add Plant profile — validation",
    "Requires type-specific fields (e.g. Make/Model for plant).",
    "Choose Plant + a client. Leave Name empty and create.",
    "Validation error on Name / Make Model. Profile is not created.",
)

step(
    22,
    "Add Plant profile — full create",
    "Creates a plant profile, uploads media, and auto-generates a QR image.",
    "Fill Make/Model, ID, Serial, Description, Notes. Optionally upload logo, pictures, documents; "
    "add a weblink and contact; add checklist items. Create.",
    "Success. Profile appears in list. View shows QR image. Media relation tabs show uploads on Edit.",
)

step(
    23,
    "Add CustomQR — live preview",
    "Creates a QR that encodes a destination URL, with live preview while typing (legacy realtimeqr).",
    "Add Profile → type CustomQR → client. Type https://example.com/test in Url. Watch QR preview update. Save.",
    "Preview image appears before save. After save, View shows a QR. Download QR works.",
)

step(
    24,
    "Add Code profile",
    "Creates a destination/bridge code profile (application + destination URL + optional bridge graphic + colour).",
    "Add Profile → type Code. Fill Application, Destination URL, Popup message. "
    "Toggle Activate Bridge Graphic. Set Colour Selector if shown. Upload logo as bridge graphic. Create.",
    "Profile saved. QR generated. Colour stored. Logo appears under Company logo on Edit.",
)

step(
    25,
    "View profile",
    "Read-only detail of profile fields, media, contacts, and QR (legacy product/view).",
    "From list, click View on a profile.",
    "Infolist sections show Profile, Videos, Web links, Pictures, Documents, Contacts, QR, "
    "Data collection, Code & security.",
)

step(
    26,
    "Download QR (PNG)",
    "Downloads the QR/Data Matrix PNG for printing or sharing.",
    "On View or Edit, click Download QR.",
    "Browser downloads a PNG (e.g. CSQRIMG{id}.png).",
)

step(
    27,
    "Download PDF",
    "Builds a PDF with profile number, name, URL, and QR image (legacy downloadQrpdf).",
    "On View or Edit, click Download PDF.",
    "Browser downloads code-{id}.pdf. Open it: text + QR image present.",
)

step(
    28,
    "QR colour change",
    "Regenerates the QR using a foreground colour (legacy qrCodeColor / colour selector).",
    "On a Code profile View/Edit, click QR colour. Pick a red/blue colour. Confirm.",
    "Success toast. QR preview/image updates. Re-download PNG and confirm colour changed.",
)

step(
    29,
    "Renew code (single)",
    "Extends code profile expiry by 1 year and creates a Renew code order.",
    "On a Code profile, click Renew code and confirm.",
    "Notification shows order id and new expiry. Manage Code Orders shows a new order with status Renew. "
    "Profile expired_at is about 1 year out.",
)

step(
    30,
    "Renew selected codes (bulk)",
    "Renews multiple code profiles in one action from the list.",
    "On Profiles list, filter Type = Code if possible. Select 2+ code profiles. "
    "Use bulk action Renew selected codes. Confirm.",
    'Success notification with count and order #. Non-code selections alone show '
    '"Please select code to be renew."',
)

step(
    31,
    "Remove logo / bridge graphic",
    "Deletes the company/bridge logo file and DB row (legacy removelogo).",
    "Edit a profile that has a logo. In Company logo relation, click Remove logo and confirm.",
    "Logo disappears from the table and from View. Placeholder / no logo state shows.",
)

step(
    32,
    "Edit profile media (pictures / docs / videos)",
    "Maintains gallery of pictures, documents, and YouTube/video links on a profile.",
    "Edit profile. Add/remove a picture, document, and video (YouTube URL or ID). Save or use relation actions.",
    "Items appear on View. Deleting an item removes it. "
    "(File upload to YouTube needs OAuth — linking by URL is enough for this test.)",
)

step(
    33,
    "Archive profile",
    "Soft-deletes a profile so it leaves the manage list without a hard DB wipe.",
    "Edit a throwaway profile → Archive (or Delete on list) and confirm.",
    "Profile no longer appears in Manage Profiles list.",
)

# ORDERS
h1("9. Orders")

step(
    34,
    "Manage Code Orders — list",
    "Lists code purchase / renew orders (legacy codeorder).",
    "Open Order → Manage Code Orders.",
    "Table shows orders with status filter (All, New, Renew, Invoice Send, Paid). "
    "Create route is not available (404) — orders come from portal/renew flows.",
)

step(
    35,
    "Manage Code Orders — view & change status",
    "Inspects order detail and updates workflow status.",
    "Open an order. Review amounts, client, line details. Use Change status (e.g. Invoice Send → Paid).",
    "Status updates and is reflected in the list with correct colour/label.",
)

step(
    36,
    "Manage Order (physical labels) — list/view",
    "Lists physical label orders (legacy order).",
    "Open Manage Order. Open one order View.",
    "Detail shows client, postage/amounts, profile link if any. "
    "Change status among New / Paid / Shipped / Completed / Cancelled as allowed.",
)

step(
    37,
    "Manage Form Builder Orders — list/view",
    "Lists form-builder purchase orders (legacy formbuilderorder). "
    "Does not include the form question editor.",
    "Open Manage Form Builder Orders. View one order.",
    "Detail loads. Status can be changed. No create screen (create 404).",
)

# SETTINGS
h1("10. Settings")

step(
    38,
    "Global Settings",
    "Edits key/value platform settings (PayPal, contact email, YouTube keys, etc.).",
    "Open Settings → Global Settings. Change a harmless field (e.g. contact email) and Save. "
    "Do not paste fake YouTube client ids if validation blocks placeholders.",
    "Success notification. Reload page — value persisted. "
    "Invalid placeholder YouTube client id is rejected if you try one.",
)

step(
    39,
    "Code Pricing",
    "Maintains retail/reseller quantity tiers used for code purchase and renew amounts (code_prising).",
    "Open Code Pricing. Edit an amount for a qty band. Save.",
    "Values persist. Renew flow uses these tiers for per-code amount.",
)

step(
    40,
    "Reseller Pricing",
    "Maintains reseller_pricing tiers for reseller clients.",
    "Open Reseller Pricing. Add or edit a qty/amount row. Save.",
    "Data persists after reload.",
)

# CMS
h1("11. Website CMS")

step(
    41,
    "Testimonials — list/create/edit",
    "Manages public testimonial content (text/video).",
    "Open Manage Testimonial. Create with required fields. Edit and View. Optionally delete a test entry.",
    "CRUD works. Empty create shows validation errors.",
)

step(
    42,
    "Gallery — add / block / delete",
    "Manages gallery images (approve/block toggle). There is no edit page (matches legacy).",
    "Open Manage Gallery. Upload one or more images. Toggle Block/Unblock. Delete a test image.",
    "Thumbnails show. Block flips approve flag. Delete removes image and files.",
)

# ACCOUNT
h1("12. Account")

step(
    43,
    "Change Password",
    "Lets the logged-in admin set a new password.",
    "Open Change Password. Enter current password Admin@12345, then a temporary new password, save. "
    "Sign out and sign in with the new password. Change it back to Admin@12345.",
    "Password change succeeds; login works with the new password. After restore, original password works again.",
)

step(
    44,
    "Sign out",
    "Ends the admin session.",
    "Use the avatar/user menu → Sign out.",
    "Redirect to login. Visiting /admin/clients redirects to login.",
)

# ROLES
h1("13. Role checks (Support vs Super Admin)")

step(
    45,
    "Support cannot write settings/clients",
    "Support role can view many areas but cannot perform privileged writes (AdminRole policies).",
    "Log in as a Support user (or temporarily set a user’s admin_role to Support). "
    "Try Global Settings and Create Client.",
    "Global Settings forbidden or hidden. Creating clients is blocked. Profiles list may still be viewable.",
)

step(
    46,
    "Super Admin full access",
    "Super Admin retains full write access.",
    "Log back in as admin@scanlink.com. Confirm Global Settings and Add Client are available.",
    "Both screens load and allow save/create.",
)

# SMOKE
h1("14. Quick smoke checklist (15 minutes)")
para("If you only have a short window, run these in order:")
for i, t in enumerate(
    [
        "Login as Super Admin",
        "Click all 12 Admin Home tiles (spot-check load)",
        "Create Client (valid data)",
        "Open Manage User for that client",
        "Add one Plant profile + View QR",
        "Add one CustomQR (confirm live preview)",
        "Add one Code profile → Download PDF + Renew code",
        "Open each of the three order lists",
        "Open Global Settings, Code Pricing, Gallery, Testimonial",
        "Change Password page loads",
        "Sign out",
    ],
    1,
):
    bullet(f"{i}. {t}")

h1("15. Defect log template")
para("Copy rows as needed while testing:")

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, name in enumerate(["Step #", "Area", "What happened", "Expected", "Severity"]):
    hdr[i].text = name
for _ in range(8):
    row = table.add_row().cells
    for c in row:
        c.text = ""

doc.add_paragraph()
h1("16. Sign-off")
para("Tester name: _______________________________")
para("Date: _______________    Environment: http://localhost:8000/admin")
para("Build / branch: _______________________________")
para("Overall result:  [ ] Pass    [ ] Pass with minor issues    [ ] Fail")
para("Comments:")
para("____________________________________________________________________")
para("____________________________________________________________________")

h1("17. Out of scope reminder")
bullet("Client portal (application/) login and dashboards")
bullet("Public QR destination / bridge pages for scanners")
bullet("Form builder question/answer editor")
bullet("Mobile apps / API consumers")
bullet("Production DB cutover and live YouTube OAuth secrets")

out1 = r"D:\projects\product\scanlink-laravel\docs\ScanLink_Admin_Manual_Test_Guide.docx"
out2 = r"D:\projects\scanlink\ScanLink_Admin_Manual_Test_Guide.docx"
doc.save(out1)
doc.save(out2)
print(out1)
print(out2)
